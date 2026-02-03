
import docx
import json
import os

def read_docx(file_path):
    doc = docx.Document(file_path)
    content = []
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text.strip())
    
    # Tables often contain package details
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
        
    return {"paragraphs": content, "tables": tables}


file_path = r"c:\Users\Admin\Downloads\Horizon\assets\docs\Tour Package for Website New Zealand.docx"
output_path = r"c:\Users\Admin\Downloads\Horizon\extracted_content.json"
if os.path.exists(file_path):
    data = read_docx(file_path)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)
    print(f"Extraction complete. Saved to {output_path}")
else:
    print(f"File not found: {file_path}")

